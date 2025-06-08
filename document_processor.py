import PyPDF2
import docx
import logging
import re
from typing import Optional, Dict, Any
from ai_analyzer import AIAnalyzer

class DocumentProcessor:
    """Handles text extraction and analysis from various document formats"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.legal_indicators = [
            r'(?i)(agreement|contract|deed|memorandum|notice|affidavit|petition|complaint)',
            r'(?i)(court|tribunal|jurisdiction|legal|law|statute|regulation)',
            r'(?i)(party|parties|hereby|whereas|witnesseth|undertaking)',
            r'(?i)(rights|obligations|liability|indemnity|warranty|termination)',
            r'(?i)(clause|section|article|provision|term|condition)',
            r'(?i)(executed|signed|dated|witness|notary|stamp|seal)',
            r'(?i)(plaintiff|defendant|petitioner|respondent|appellant)',
            r'(?i)(compensation|damages|penalty|breach|default|remedy)',
            r'(?i)(arbitration|mediation|dispute|settlement|resolution)',
            r'(?i)(license|permit|authorization|compliance|enforcement)'
        ]

    def is_legal_document(self, text: str) -> bool:
        """Check if the document is legal-related"""
        # Count how many legal indicators are present
        indicator_count = sum(1 for pattern in self.legal_indicators if re.search(pattern, text))
        
        # If more than 2 legal indicators are present, consider it a legal document
        return indicator_count > 2

    def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process the document and return analysis results"""
        try:
            # Extract text from document
            text = self.extract_text(file_path, file_type)
            
            # Get basic document stats
            stats = self.get_document_stats(text)
            
            # Check if it's a legal document
            if not self.is_legal_document(text):
                return {
                    'is_legal': False,
                    'message': 'This document does not appear to be a legal document. I can only provide detailed analysis for legal documents such as contracts, agreements, notices, court documents, etc.',
                    'summary': self.generate_basic_summary(text),
                    'stats': stats
                }
            
            # Process legal document
            analyzer = AIAnalyzer()
            analysis_results = analyzer.analyze_document(text, file_type)
            
            return {
                'is_legal': True,
                'analysis': analysis_results,
                'summary': analysis_results.get('summary', ''),
                'stats': stats
            }
            
        except Exception as e:
            self.logger.error(f"Document processing error: {str(e)}")
            return {
                'error': 'Error processing document',
                'message': str(e)
            }

    def generate_basic_summary(self, text: str) -> str:
        """Generate a basic summary for non-legal documents"""
        try:
            # Split into sentences (handle common abbreviations)
            text = text.replace('Mr.', 'Mr').replace('Mrs.', 'Mrs').replace('Dr.', 'Dr')
            sentences = [s.strip() + '.' for s in text.split('.') if s.strip()]
            
            if len(sentences) <= 3:
                return text
            
            # For longer documents, take first 2 and last sentence
            summary = ' '.join([
                sentences[0],
                sentences[1],
                '...',
                sentences[-1]
            ])
            
            return summary
            
        except Exception as e:
            self.logger.error(f"Summary generation error: {str(e)}")
            return "Could not generate summary due to an error."

    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from document based on file type"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type.lower() == 'docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text
                    except Exception as e:
                        self.logger.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                        
        except Exception as e:
            self.logger.error(f"Error reading PDF file: {str(e)}")
            raise
        
        return self._clean_text(text)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
                        
        except Exception as e:
            self.logger.error(f"Error reading DOCX file: {str(e)}")
            raise
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Join lines back
        cleaned_text = '\n'.join(lines)
        
        return cleaned_text.strip()
    
    def get_document_stats(self, text: str) -> dict:
        """Get basic statistics about the document"""
        if not text:
            return {
                'word_count': 0,
                'character_count': 0,
                'paragraph_count': 0,
                'line_count': 0
            }
        
        words = text.split()
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        lines = [l for l in text.split('\n') if l.strip()]
        
        return {
            'word_count': len(words),
            'character_count': len(text),
            'paragraph_count': len(paragraphs),
            'line_count': len(lines)
        }
